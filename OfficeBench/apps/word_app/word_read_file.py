import os
import fire
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# DEMO = (
#     'You can create a new word file by calling `word_read_file` with 1 arguments.\n'
#     '1. The path of the word file: file_path: str\n'
#     "You can call it by: {'app': 'word', 'action': 'word_read_file', 'file_path': ...}"
# )

DEMO = (
    'read the content of the word file: '
    '{"app": "word", "action": "read_file", "file_path": [THE_PATH_TO_THE_WORD_FILE]}'
)


def construct_action(work_dir, args: dict, py_file_path='/apps/word_app/word_read_file.py'):
    # TODO: not sure if we need to specify the file path with the current workdir
    return f'python3 {py_file_path} --file_path {args["file_path"]}'

def word_read_file(file_path):
    doc = Document(file_path)
    return doc.paragraphs

def word_table_to_string(table):
    def iter_cell_contents(c):
        for content in c.iter_inner_content():
            if isinstance(content, Paragraph):
                yield content.text
            elif isinstance(content, Table):
                yield word_table_to_string(content)
            else:
                raise ValueError(f"Unsupported content type: {type(content)}")
    def cell_to_string(c):
        row = "|"+"|".join(iter_cell_contents(c))+"|"
        return row
    
    def iter_row_cell_texts(row):
        for c in row.cells:
            yield cell_to_string(c)

    def row_to_string(row):
        row = "|"+ "|".join(iter_row_cell_texts(row))+"|"
        return row
    return "\n".join(
        [row_to_string(r) for r in table.rows]
    )


def word_read_file_into_string(file_path):
    doc = Document(file_path)
    txt_builder = []
    for c in doc.iter_inner_content():
        # print(type(c), c.text if isinstance(c, Paragraph) else "")
        if isinstance(c, Paragraph):
            txt_builder.append(c.text)
        elif isinstance(c, Table):
            txt_builder.append(word_table_to_string(c))
        else:
            raise ValueError(f"Unsupported content type: {type(c)}")
    return "\n".join(txt_builder)
    string = '\n'.join([paragraph.text for paragraph in doc.iter_inner_content()])
    return string
        
def wrap(paragraphs):
    ob = ("The following is the content from the word file:")
    for paragraph in paragraphs:
        ob += f"\n{paragraph.text}"
    return ob

def main(file_path, debug=False):
    if not os.path.exists(file_path):
        return f"OBSERVATION: The file {file_path} does not exist. Failed to read the file."
    #paragraphs = word_read_file(file_path)
    content = word_read_file_into_string(file_path)
    if debug:
        print(content)
    #observation = wrap(content)
    return 'OBSERVATION: ' + content


if __name__ == '__main__':
    fire.Fire(main)